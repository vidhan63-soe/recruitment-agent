"""
Resume Parser Service.
Extracts text from PDF/DOCX, detects contact info, and identifies sections.
Optimized for speed — no heavy NLP models, uses regex + heuristics.
"""

import re
import uuid
from pathlib import Path
from io import BytesIO

import pdfplumber
from docx import Document
from loguru import logger


# ── Contact Info Patterns ──────────────────────

EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_PATTERN = re.compile(
    r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}"
)

# Section header patterns (case-insensitive)
SECTION_PATTERNS = {
    "education": re.compile(
        r"^(?:education|academic|qualification|degree)", re.IGNORECASE | re.MULTILINE
    ),
    "experience": re.compile(
        r"^(?:experience|employment|work\s*history|professional\s*experience|career)",
        re.IGNORECASE | re.MULTILINE,
    ),
    "skills": re.compile(
        r"^(?:skills|technical\s*skills|core\s*competenc|proficienc|expertise|technologies)",
        re.IGNORECASE | re.MULTILINE,
    ),
    "projects": re.compile(
        r"^(?:projects|personal\s*projects|key\s*projects)",
        re.IGNORECASE | re.MULTILINE,
    ),
    "certifications": re.compile(
        r"^(?:certif|licens|accreditation)", re.IGNORECASE | re.MULTILINE
    ),
    "summary": re.compile(
        r"^(?:summary|objective|profile|about\s*me|professional\s*summary)",
        re.IGNORECASE | re.MULTILINE,
    ),
}


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Also try extracting tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            cleaned = [str(cell).strip() for cell in row if cell]
                            if cleaned:
                                text_parts.append(" | ".join(cleaned))
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Could not parse PDF: {e}")

    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        doc = Document(BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Could not parse DOCX: {e}")


def extract_contact_info(text: str) -> dict:
    """Pull out name, email, phone from resume text."""
    info = {"name": "Unknown", "email": "", "phone": ""}

    # Email
    email_match = EMAIL_PATTERN.search(text)
    if email_match:
        info["email"] = email_match.group()

    # Phone
    phone_match = PHONE_PATTERN.search(text)
    if phone_match:
        info["phone"] = phone_match.group().strip()

    # Name heuristic: first non-empty line that isn't an email/phone/url
    lines = text.strip().split("\n")
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if not line or len(line) < 2 or len(line) > 60:
            continue
        if "@" in line or "http" in line or line.startswith("+"):
            continue
        if re.match(r"^[\d\s\-()+]+$", line):  # Pure phone number
            continue
        # Likely a name
        info["name"] = line.title()
        break

    return info


def extract_sections(text: str) -> dict[str, str]:
    """
    Identify resume sections by header keywords.
    Returns dict mapping section name to its content.
    """
    sections = {}
    lines = text.split("\n")

    # Find all section headers and their line positions
    found_sections = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if len(stripped) < 3 or len(stripped) > 60:
            continue
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(stripped):
                found_sections.append((i, section_name))
                break

    # Extract content between section headers
    for idx, (line_num, section_name) in enumerate(found_sections):
        start = line_num + 1
        end = found_sections[idx + 1][0] if idx + 1 < len(found_sections) else len(lines)
        content = "\n".join(lines[start:end]).strip()
        if content:
            sections[section_name] = content

    return sections


def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> list[str]:
    """
    Split text into overlapping chunks by character count.
    Tries to break at sentence boundaries for cleaner chunks.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at a sentence boundary
        if end < len(text):
            # Look for sentence-ending punctuation near the end
            search_zone = text[max(end - 80, start):end]
            last_period = max(
                search_zone.rfind(". "),
                search_zone.rfind(".\n"),
                search_zone.rfind("? "),
                search_zone.rfind("! "),
            )
            if last_period > 0:
                end = max(end - 80, start) + last_period + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap

    return chunks


def parse_resume(
    file_bytes: bytes,
    filename: str,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> dict:
    """
    Full parsing pipeline for a single resume.
    Returns a dict ready to be converted to ParsedResume.
    """
    ext = Path(filename).suffix.lower()
    resume_id = str(uuid.uuid4())[:12]

    logger.info(f"Parsing resume: {filename} (id: {resume_id})")

    # 1. Extract raw text
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported format: {ext}")

    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError(f"Resume appears empty or too short: {filename}")

    # 2. Extract contact info
    contact = extract_contact_info(raw_text)

    # 3. Identify sections
    sections = extract_sections(raw_text)

    # 4. Chunk text for embeddings
    text_chunks = chunk_text(raw_text, chunk_size, chunk_overlap)

    # Build chunk objects
    chunks = []
    for i, chunk_text_content in enumerate(text_chunks):
        chunks.append({
            "chunk_id": f"{resume_id}_chunk_{i}",
            "resume_id": resume_id,
            "text": chunk_text_content,
            "chunk_index": i,
            "metadata": {
                "filename": filename,
                "candidate_name": contact["name"],
                "section": _identify_chunk_section(chunk_text_content, sections),
            },
        })

    logger.info(
        f"Parsed {filename}: name={contact['name']}, "
        f"sections={list(sections.keys())}, chunks={len(chunks)}"
    )

    return {
        "resume_id": resume_id,
        "filename": filename,
        "candidate_name": contact["name"],
        "email": contact["email"],
        "phone": contact["phone"],
        "raw_text": raw_text,
        "sections": sections,
        "chunks": chunks,
    }


def _identify_chunk_section(chunk_text: str, sections: dict) -> str:
    """Determine which section a chunk most likely belongs to."""
    best_match = "general"
    best_overlap = 0

    for section_name, section_content in sections.items():
        # Simple overlap check — how many words in common
        chunk_words = set(chunk_text.lower().split())
        section_words = set(section_content.lower().split()[:100])
        overlap = len(chunk_words & section_words)
        if overlap > best_overlap:
            best_overlap = overlap
            best_match = section_name

    return best_match
